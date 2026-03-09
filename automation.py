#pip install python-docx -> to install the document packages for python.
#pip install deep_translator -> to install the translation packages

from docx import Document #imports the Document class from the python-docx library, which is used to create and manipulate Word documents in Python.
import requests #imports the requests library, which is used to make HTTP requests in Python.
import os
import platform
from datetime import datetime
from deep_translator import GoogleTranslator

now = datetime.now()
todays_date = now.strftime("%d:%m:%Y")

translator=GoogleTranslator(source='en',target='es')

def cve_search(cve_id):
  response=requests.get(f"https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}")

  if response.status_code == 200:
    data=response.json()
    # Check if 'vulnerabilities' key exists and is not empty
    if not data.get("vulnerabilities"):
        return None, None, None # No vulnerabilities found for this CVE ID

    cve_data = data["vulnerabilities"][0]["cve"] #gives the first item in the list
    description = cve_data["descriptions"][0]["value"]
    references = cve_data["references"]

    cvss_score = None # Initialize cvss_score

    metrics = cve_data.get("metrics")
    if metrics:
        if "cvssMetricV31" in metrics and metrics["cvssMetricV31"]:
            cvss_score = metrics["cvssMetricV31"][0]["cvssData"]["baseScore"]
        elif "cvssMetricV2" in metrics and metrics["cvssMetricV2"]:
            # Fallback to CVSS v2 if V3.1 is not available
            cvss_score = metrics["cvssMetricV2"][0]["cvssData"]["baseScore"]

    return description,references,cvss_score
  else:
    return None,None,None

cve_id=input("Provide a CVE-ID:").strip()
description,references,cvss_score=cve_search(cve_id)
result=""

if cvss_score is None:
  result = "Unknown"
elif cvss_score < 4.0: # CVSS v3.1: 0.1-3.9 Low
  result = "Low"
elif cvss_score < 7.0: # CVSS v3.1: 4.0-6.9 Medium
  result = "Medium"
elif cvss_score < 9.0: # CVSS v3.1: 7.0-8.9 High
  result = "High"
else: # CVSS v3.1: 9.0-10.0 Critical
  result = "Critical"

dic={"Date Created": todays_date,
    "CVE Detail":cve_id,
    "Description":description,
     "References":references,
     "Metric":[cvss_score,result]}

severity_map = {
    "Low": "Bajo",
    "Medium": "Medio",
    "High": "Alto",
    "Critical": "Crítico",
    "Unknown": "Desconocido"
}

resultado_es = severity_map.get(result, "Desconocido")

docx=Document() #Initialises the document
docx.add_heading("CVE SECURITY REPORT",level=1)
docx.add_heading(dic["CVE Detail"],1)

docx.add_heading("Date Created",2)
docx.add_paragraph(dic["Date Created"])

docx.add_heading("Description",2)
docx.add_paragraph(dic["Description"])

docx.add_heading("References",2)
for item in dic["References"]:
  docx.add_paragraph(str(item),style="List Bullet") #Put the listed items into a bullet points
  

docx.add_heading("Metric",2)
metric_line=f"{cvss_score} ({result})"
docx.add_paragraph(metric_line,style="List Bullet") 

docx.add_page_break() #creates a page break seperating the En version from the Es version, making for a better legibility
docx.add_heading("INFORME DE SEGURIDAD DE CVE")
docx.add_heading("Fecha Creada",2)
docx.add_paragraph(dic["Date Created"])

docx.add_heading("Descripción",level=2)
docx.add_paragraph(translator.translate(dic["Description"])) 

docx.add_heading("Referencias",2)
for item in dic["References"]:
  docx.add_paragraph(str(item),style="List Bullet")
  

docx.add_heading("Métricas",2)
metric_line=f"{cvss_score} ({resultado_es})"
docx.add_paragraph(metric_line,style="List Bullet")

docx.save(f"{cve_id}.docx")
print("The CVE document has been initialised and saved!")

if platform.system() == "Windows":
  os.startfile(f"{cve_id}.docx") #This opens the document created
else:
  os.system(f"open {cve_id}.docx")
